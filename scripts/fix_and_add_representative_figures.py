from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测_格式修订版.docx"
FIG_DIR = PROJECT_ROOT / "paper_figures"


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_p._p)
    if style:
        new_p.style = style
    if text:
        new_p.add_run(text)
    return new_p


def add_picture_block(anchor, image_name: str, caption: str, source: str, intro: str) -> None:
    current = insert_after(anchor, intro, "Normal")
    img_p = insert_after(current, style="Normal")
    img_p.add_run().add_picture(str(FIG_DIR / image_name), width=Cm(14.2))
    caption_p = insert_after(img_p, caption, "Caption")
    insert_after(caption_p, source, "annotation text")


def main() -> None:
    doc = Document(DOC_PATH)

    # 先修正重复图号：后面的“实验验证思路图”改成图5
    for p in doc.paragraphs:
        if p.text.strip() == "图4. 实验验证思路图":
            p.text = "图5. 实验验证思路图"
            break

    paragraphs = {p.text.strip(): p for p in doc.paragraphs}

    # 在模型对比实验后加入一张模型对比图
    p = paragraphs.get("尽管如此，该实验仍具有一定参考价值。首先，规则基线可作为系统最基础的比较对象；其次，逻辑回归和随机森林表明人工设计特征能够支持传统监督分类建模；最后，对比实验说明在当前特征体系下，随机森林更适合作为后续严格实验的主模型。为了降低“标签规则与模型特征同源”带来的影响，后文还进一步给出特征消融实验和更严格的按题目划分泛化实验。")
    if p is not None:
        add_picture_block(
            p,
            "fig_model_comparison.png",
            "图6. 模型对比实验结果可视化",
            "数据来源：本系统模型对比实验结果整理",
            "为了更直观地展示不同方法在主要评价指标上的表现，本文进一步将规则基线、逻辑回归和随机森林的实验结果绘制为柱状图，如图6所示。",
        )

    # 在特征消融实验后加入消融图
    paragraphs = {p.text.strip(): p for p in doc.paragraphs}
    p = paragraphs.get("实验结果显示，完整特征集对应的模型表现最好；当去除 base_similarity 后，模型性能仅出现较小幅度下降，说明除基础相似度外，Token 序列、词法统计和结构差异等特征同样能够提供有效判别信息。与此同时，仅使用 base_similarity 也能取得较高结果，说明规则相似度本身已经具备较强表达能力。综合来看，本文所设计的多维特征体系具有一定冗余性和互补性，在弱监督环境下能够保持较稳定效果。")
    if p is not None:
        add_picture_block(
            p,
            "fig_ablation_f1.png",
            "图7. 特征消融实验 F1 值对比",
            "数据来源：本系统特征消融实验结果整理",
            "除表格形式外，本文还对不同特征组合下的 F1 值进行了可视化展示，从而更直观地反映各类特征对最终性能的影响，如图7所示。",
        )

    # 在严格泛化实验后加入严格实验图
    paragraphs = {p.text.strip(): p for p in doc.paragraphs}
    p = paragraphs.get("实验结果表明，在 68 个训练题目和 30 个测试题目条件下，随机森林模型在 48152 个训练样本和 19947 个测试样本上取得了较高性能，其中准确率为 0.99995，精确率为 0.99913，召回率为 1.00000，F1 值为 0.99957，ROC-AUC 为 0.99999。该结果说明，本文设计的特征不仅能够在同题环境下工作，也能够在一定程度上迁移到未参与训练的新题目上，从而增强了方法的泛化可信度。")
    if p is not None:
        add_picture_block(
            p,
            "fig_question_split_metrics.png",
            "图8. 按题目划分严格泛化实验结果",
            "数据来源：本系统严格泛化实验结果整理",
            "为了进一步说明模型在更严格设置下的表现，本文将按题目划分泛化实验的主要指标绘制为图8。",
        )

    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()
