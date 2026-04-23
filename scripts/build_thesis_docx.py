from __future__ import annotations

import json
import shutil
import tempfile
import zipfile
from pathlib import Path
import re

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PROJECT_ROOT / "2.1工科论文参考模板.docx"
DRAFT_PATH = PROJECT_ROOT / "thesis_draft.md"
OUTPUT_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测.docx"
PAPER_TABLES_DIR = PROJECT_ROOT / "paper_tables"
ABLATION_TABLE_PATH = PROJECT_ROOT / "outputs_ablation" / "ablation_metrics.csv"
QUESTION_SPLIT_PATH = PROJECT_ROOT / "outputs_question_split" / "question_split_metrics.json"


def normalize_text(raw_text: str) -> str:
    """去除 Markdown 痕迹并做轻量文本清洗。"""

    return raw_text.replace("**", "").replace("`", "").strip()


def markdown_to_sections(markdown_text: str) -> list[tuple[int, str]]:
    """把 Markdown 粗略转换成标题与正文段落序列。

    这里只处理当前草稿用到的一级、二级、三级标题和普通段落，
    目的是尽量把草稿内容按模板样式写入 Word，而不是做完整 Markdown 解析。
    """

    sections: list[tuple[int, str]] = []
    for raw_line in markdown_text.splitlines():
        line = normalize_text(raw_line)
        if not line:
            sections.append((0, ""))
            continue
        if line.startswith("### "):
            sections.append((3, line[4:].strip()))
        elif line.startswith("## "):
            sections.append((2, line[3:].strip()))
        elif line.startswith("# "):
            sections.append((1, line[2:].strip()))
        else:
            sections.append((0, line))
    return sections


def clear_document_body(doc: Document) -> None:
    """清空模板中的示例正文，但保留封面/节属性等模板骨架。"""

    body = doc._element.body
    children = list(body)
    preserve_prefix = 2 if len(children) >= 2 else 0
    for child in children[preserve_prefix:-1]:
        body.remove(child)


def add_paragraph_with_style(doc: Document, text: str, style_name: str) -> None:
    paragraph = doc.add_paragraph(style=style_name)
    if text:
        paragraph.add_run(text)


def preferred_body_style(doc: Document) -> str:
    """优先使用与模板示例一致的正文样式。"""

    return "Normal"


def strip_section_number(title: str) -> str:
    """将草稿中的 1.1 / 2.3 等编号去掉，更贴近模板二级标题格式。"""

    parts = title.split(" ", 1)
    candidate = parts[0]
    if candidate.count(".") == 1 and candidate.replace(".", "").isdigit():
        return parts[1] if len(parts) > 1 else title
    return title


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """按模板正文风格设置图示单元格文本。"""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 1
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    """给 Word 表格单元格添加浅色底纹，形成可编辑示意图。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_diagram_table(doc: Document, rows: list[list[str]], caption: str, source: str, header: bool = False) -> None:
    """使用模板内可编辑表格生成论文示意图。"""

    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            set_cell_text(cell, value, bold=header and row_index == 0)
            shade_cell(cell, "D9EAF7" if row_index == 0 and header else "F4F9FC")
    add_paragraph_with_style(doc, caption, "Caption")
    add_paragraph_with_style(doc, f"数据来源：{source}", "annotation text")


def add_system_flow_figure(doc: Document) -> None:
    """插入系统整体流程图。"""

    add_diagram_table(
        doc,
        [
            ["数据读取", "代码预处理", "特征提取", "弱监督标注", "模型训练", "相似度预测", "报告与复核"],
            ["读取 Java 提交并按题目分组", "注释清理、字面量与标识符归一化", "计算词法、序列、结构和长度特征", "依据基础相似度生成伪标签", "训练逻辑回归与随机森林模型", "输出模型相似概率与高相似代码对", "生成表格、报告和人工复核样本"],
        ],
        "图1. 系统整体流程图",
        "根据本文系统设计整理",
        header=True,
    )


def add_module_figure(doc: Document) -> None:
    """插入系统功能模块图。"""

    add_diagram_table(
        doc,
        [
            ["代码相似度检测系统"],
            ["数据加载模块", "预处理模块", "特征工程模块"],
            ["样本构造模块", "模型训练模块", "预测报告模块"],
            ["对比实验模块", "严格评估模块", "人工复核模块"],
        ],
        "图2. 系统主要功能模块图",
        "根据本文系统实现整理",
    )


def add_feature_figure(doc: Document) -> None:
    """插入特征体系图。"""

    add_diagram_table(
        doc,
        [
            ["特征类别", "代表特征", "作用说明"],
            ["词法相似度", "Token Jaccard、关键字余弦、运算符余弦", "反映代码词汇和语言构造的重合程度"],
            ["序列相似度", "Token 序列相似度", "反映语句顺序和实现流程的一致性"],
            ["结构相似度", "结构 Token、方法数、循环数、分支数", "描述程序控制结构和组织方式"],
            ["规模差异", "Token 数量、代码行数", "刻画两份代码在整体规模上的接近程度"],
            ["基础相似度", "多项规则特征加权融合", "用于弱监督标签构造和结果解释"],
        ],
        "图3. 代码相似度特征体系图",
        "根据本文特征工程设计整理",
        header=True,
    )


def add_evaluation_figure(doc: Document) -> None:
    """插入实验验证思路图。"""

    add_diagram_table(
        doc,
        [
            ["验证层次", "实验内容", "主要目的"],
            ["伪标签一致性评估", "规则基线、逻辑回归、随机森林对比", "检验模型对弱监督标签体系的拟合情况"],
            ["严格泛化评估", "跨用户过滤、按题目划分、移除 base_similarity", "降低同源特征和随机切分带来的偏乐观风险"],
            ["人工辅助复核", "高分样本复核与分层抽样机制", "补充自动指标之外的人工判断依据"],
        ],
        "图4. 实验验证思路图",
        "根据本文实验设计整理",
        header=True,
    )


def add_reference_list(doc: Document) -> None:
    """添加当前论文的参考文献列表。"""

    references = [
        "[1] DSA Dataset[DB/OL]. Zenodo, 2026. Available: https://zenodo.org/records/18136250.",
        "[2] 吴贺. 前后端解耦模式及开发[J]. 计算机系统应用, 2017, 26(2): 217-221.",
        "[3] 李智慧. 大型网站技术架构[M]. 北京: 电子工业出版社, 2013.",
        "[4] 许令波. 深入分析Java Web技术内幕: 第2版[M]. 北京: 电子工业出版社, 2014.",
        "[5] Breiman L. Random Forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[6] Ratner A, Bach S, Ehrenberg H, et al. Snorkel: Rapid Training Data Creation with Weak Supervision[J]. Proceedings of the VLDB Endowment, 2017, 11(3): 269-282.",
        "[7] Jurafsky D, Martin J H. Speech and Language Processing[M]. 3rd draft ed. 2023.",
        "[8] Manning C D, Raghavan P, Schutze H. Introduction to Information Retrieval[M]. Cambridge: Cambridge University Press, 2008.",
        "[9] 王斌, 张海军. 程序代码相似度检测研究综述[J]. 计算机工程与应用, 相关年份后续可按学校数据库检索补全.",
        "[10] 弱监督学习、代码相似度检测及程序查重相关中文文献，后续可继续按学校图书馆数据库补充与替换。",
    ]
    for item in references:
        add_paragraph_with_style(doc, item, "Normal")


def add_table_from_dataframe(doc: Document, dataframe: pd.DataFrame, caption: str, source: str) -> None:
    """把真实实验数据表插入到论文中。"""

    if dataframe.empty:
        return
    add_paragraph_with_style(doc, caption, "Caption")
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(dataframe.columns):
        table.rows[0].cells[idx].text = str(column)
    for row in dataframe.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    add_paragraph_with_style(doc, f"数据来源：{source}", "annotation text")


def add_toc_field(paragraph) -> None:
    """插入 Word 目录域，打开文档后可更新目录。"""

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Word 中右键更新域后显示。"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def insert_tables_if_needed(doc: Document, text: str) -> None:
    """在对应段落后插入论文实验表。"""

    if text.endswith("系统整体流程如图1所示。"):
        add_system_flow_figure(doc)
    elif text.endswith("系统主要功能模块如图2所示。"):
        add_module_figure(doc)
    elif text.endswith("本文所构建的特征体系如图3所示。"):
        add_feature_figure(doc)
    elif text.endswith("整体实验验证思路如图4所示。"):
        add_evaluation_figure(doc)
    elif text.endswith("题目提交数量分布见表1，训练样本标签分布见表2。"):
        question_table = PAPER_TABLES_DIR / "table_question_distribution.csv"
        label_table = PAPER_TABLES_DIR / "table_label_distribution.csv"
        if question_table.exists():
            df = pd.read_csv(question_table).head(10)
            add_table_from_dataframe(doc, df, "表1. 题目提交数量分布（前10项）", "Zenodo DSA Dataset 统计结果")
        if label_table.exists():
            df = pd.read_csv(label_table)
            add_table_from_dataframe(doc, df, "表2. 训练样本标签分布", "本系统跨用户训练样本统计")
    elif "表3" in text and "模型对比" in text:
        compare_table = PAPER_TABLES_DIR / "table_model_comparison.csv"
        if compare_table.exists():
            df = pd.read_csv(compare_table)
            columns = [column for column in ["method", "accuracy", "precision", "recall", "f1", "roc_auc"] if column in df.columns]
            add_table_from_dataframe(doc, df[columns], "表3. 模型对比实验结果", "本系统模型对比实验结果")
    elif text.endswith("消融实验结果见表4。"):
        if ABLATION_TABLE_PATH.exists():
            df = pd.read_csv(ABLATION_TABLE_PATH)
            columns = [column for column in ["ablation", "accuracy", "precision", "recall", "f1", "roc_auc"] if column in df.columns]
            add_table_from_dataframe(doc, df[columns], "表4. 特征消融实验结果", "本系统特征消融实验结果")
    elif text.endswith("严格泛化实验结果见表5。"):
        if QUESTION_SPLIT_PATH.exists():
            metrics = json.loads(QUESTION_SPLIT_PATH.read_text(encoding="utf-8"))
            df = pd.DataFrame(
                [
                    {
                        "model_type": metrics["model_type"],
                        "train_activities": metrics["train_activities"],
                        "test_activities": metrics["test_activities"],
                        "train_samples": metrics["train_samples"],
                        "test_samples": metrics["test_samples"],
                        "accuracy": metrics["accuracy"],
                        "precision": metrics["precision"],
                        "recall": metrics["recall"],
                        "f1": metrics["f1"],
                        "roc_auc": metrics["roc_auc"],
                    }
                ]
            )
            add_table_from_dataframe(doc, df, "表5. 按题目划分的严格泛化实验结果", "本系统严格泛化实验结果")
    elif text.endswith("人工辅助复核统计结果见表6。"):
        review_table = PAPER_TABLES_DIR / "table_manual_review_summary.csv"
        if review_table.exists():
            df = pd.read_csv(review_table)
            add_table_from_dataframe(doc, df, "表6. 高分样本人工辅助复核统计结果", "本系统高分样本人工辅助复核结果")
    elif "人工辅助复核统计结果见表6。" in text:
        review_table = PAPER_TABLES_DIR / "table_manual_review_summary.csv"
        if review_table.exists():
            df = pd.read_csv(review_table)
            add_table_from_dataframe(doc, df, "表6. 高分样本人工辅助复核统计结果", "本系统高分样本人工辅助复核结果")


def restore_template_section_properties(template_path: Path, output_path: Path) -> None:
    """把模板中的节属性、页眉页脚引用恢复到输出文档。

    python-docx 在新增节或重写正文时，可能改写 `sectPr` 中的
    `headerReference/footerReference/type` 等配置，导致页眉页脚与模板
    不完全一致。这里直接把模板中的 `sectPr` 结构覆盖到输出文档，
    尽量保证最终版式与模板保持一致。
    """

    sect_pattern = re.compile(r"(<w:sectPr[\s\S]*?</w:sectPr>)")
    with zipfile.ZipFile(template_path, "r") as template_zip:
        template_document_xml = template_zip.read("word/document.xml").decode("utf-8")
        template_sect_blocks = sect_pattern.findall(template_document_xml)

    with zipfile.ZipFile(output_path, "r") as output_zip:
        file_map = {name: output_zip.read(name) for name in output_zip.namelist()}
    output_document_xml = file_map["word/document.xml"].decode("utf-8")
    output_matches = list(sect_pattern.finditer(output_document_xml))

    if not template_sect_blocks or not output_matches:
        return

    replace_count = min(len(template_sect_blocks), len(output_matches))
    rebuilt_parts: list[str] = []
    last_index = 0
    for match_index, match in enumerate(output_matches):
        rebuilt_parts.append(output_document_xml[last_index : match.start()])
        if match_index < replace_count:
            rebuilt_parts.append(template_sect_blocks[match_index])
        else:
            rebuilt_parts.append(match.group(0))
        last_index = match.end()
    rebuilt_parts.append(output_document_xml[last_index:])
    file_map["word/document.xml"] = "".join(rebuilt_parts).encode("utf-8")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_path = Path(tmp_file.name)
    with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as new_zip:
        for name, data in file_map.items():
            new_zip.writestr(name, data)
    shutil.move(str(temp_path), output_path)


def build_document() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"找不到模板文件: {TEMPLATE_PATH}")
    if not DRAFT_PATH.exists():
        raise FileNotFoundError(f"找不到论文草稿文件: {DRAFT_PATH}")

    shutil.copyfile(TEMPLATE_PATH, OUTPUT_PATH)
    doc = Document(OUTPUT_PATH)
    clear_document_body(doc)

    draft_text = DRAFT_PATH.read_text(encoding="utf-8")
    sections = markdown_to_sections(draft_text)

    title_cn = "基于弱监督训练建模的本科代码相似度检测算法设计与实现"
    title_en = "Design and Implementation of Undergraduate Code Similarity Detection Based on Weakly Supervised Training"

    filtered_sections = [(level, text) for level, text in sections if text not in {title_cn, title_en}]

    chinese_abstract_parts: list[str] = []
    chinese_keywords = ""
    english_abstract_parts: list[str] = []
    english_keywords = ""
    body_sections: list[tuple[int, str]] = []

    mode = "before_abstract"
    for level, text in filtered_sections:
        if level == 2 and text == "摘要":
            mode = "chinese_abstract"
            continue
        if level == 2 and text == "Abstract":
            mode = "english_abstract"
            continue
        if level == 2 and text.startswith("第") and "章" in text:
            mode = "body"
        if mode == "chinese_abstract":
            if not text:
                continue
            if text.startswith("关键词"):
                chinese_keywords = text
            else:
                chinese_abstract_parts.append(text)
            continue
        if mode == "english_abstract":
            if not text:
                continue
            if text.startswith("Keywords"):
                english_keywords = text
            else:
                english_abstract_parts.append(text)
            continue
        if mode == "body":
            body_sections.append((level, text))

    for _ in range(2):
        doc.add_paragraph("")
    add_paragraph_with_style(doc, title_cn, "摘要、参考文献、注释")
    add_paragraph_with_style(doc, f"摘  要：{' '.join(chinese_abstract_parts)}", "Normal")
    add_paragraph_with_style(doc, chinese_keywords, "Normal")
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_paragraph_with_style(doc, title_en, "摘要、参考文献、注释")
    add_paragraph_with_style(doc, f"Abstract: {' '.join(english_abstract_parts)}", "Normal")
    add_paragraph_with_style(doc, english_keywords, "Normal")
    for _ in range(3):
        doc.add_paragraph("")

    # 新节开始正文，尽量贴近模板的双节结构
    doc.add_section(WD_SECTION.NEW_PAGE)

    chapter_index = 0
    in_body = False
    for level, text in body_sections:
        if level == 2 and text.startswith("第") and "章" in text:
            in_body = True
            chapter_index += 1
            chapter_title = text.split(" ", 1)[1] if " " in text else text
            chinese_number = ["一", "二", "三", "四", "五", "六", "七", "八", "九"][chapter_index - 1]
            add_paragraph_with_style(doc, f"{chinese_number}、{chapter_title}", "一级标题")
            continue
        if not in_body:
            continue
        if level == 3:
            add_paragraph_with_style(doc, strip_section_number(text), "二级标题")
        elif level == 1:
            continue
        else:
            if text == "参考文献（待整理）":
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                add_paragraph_with_style(doc, "参考文献", "一级标题")
                add_paragraph_with_style(
                    doc,
                    "放在新的一页，中文在前，英文在后，各自按照第一作者姓氏排列。正文中所列文献和这里列出文献要一一对应。",
                    "annotation text",
                )
                add_reference_list(doc)
                break
            if not text:
                continue
            style_name = preferred_body_style(doc)
            add_paragraph_with_style(doc, text, style_name)
            insert_tables_if_needed(doc, text)

    doc.save(OUTPUT_PATH)
    restore_template_section_properties(TEMPLATE_PATH, OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
