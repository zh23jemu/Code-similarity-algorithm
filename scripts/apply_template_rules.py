from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parent.parent
SOURCE_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测.docx"
OUTPUT_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测_格式修订版.docx"


def is_mostly_ascii(text: str) -> bool:
    """简单判断一段文字是否主要由英文/数字构成。"""

    stripped = "".join(ch for ch in text if not ch.isspace())
    if not stripped:
        return False
    ascii_count = sum(1 for ch in stripped if ord(ch) < 128)
    return ascii_count / max(len(stripped), 1) > 0.7


def set_run_font(run, east_asia: str, western: str, size_pt: float) -> None:
    """统一设置 run 字体，兼顾中文与英文显示。"""

    run.font.size = Pt(size_pt)
    run.font.name = western
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        east_asia,
    )
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
        western,
    )
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
        western,
    )


def apply_template_rules() -> Path:
    """按模板中可提取出的格式要求修正文档。"""

    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"找不到源论文文件: {SOURCE_PATH}")

    doc = Document(SOURCE_PATH)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"

        # 摘要标题与英文标题保留原样式，仅统一 run 字体。
        if style_name == "摘要、参考文献、注释":
            for run in paragraph.runs:
                if is_mostly_ascii(run.text):
                    set_run_font(run, "仿宋", "Times New Roman", 10.5)
                else:
                    set_run_font(run, "仿宋", "Times New Roman", 10.5)
            continue

        # 一级、二级标题维持模板标题样式。
        if style_name == "一级标题":
            for run in paragraph.runs:
                set_run_font(run, "黑体", "Times New Roman", 14)
            continue
        if style_name == "二级标题":
            for run in paragraph.runs:
                set_run_font(run, "黑体", "Times New Roman", 12)
            continue

        # 图表题与注释说明。
        if style_name == "Caption":
            for run in paragraph.runs:
                set_run_font(run, "宋体", "Times New Roman", 10)
            continue
        if style_name == "annotation text":
            for run in paragraph.runs:
                set_run_font(run, "仿宋", "Times New Roman", 10.5)
            continue

        # 正文和摘要内容：中文仿宋小四，英文 Times New Roman 小四。
        for run in paragraph.runs:
            if is_mostly_ascii(run.text):
                set_run_font(run, "仿宋", "Times New Roman", 12)
            else:
                set_run_font(run, "仿宋", "Times New Roman", 12)

        # 按模板规则微调几个常见标签文本。
        if text.startswith("关键词："):
            paragraph.text = text.replace("关键词：", "关键词：")
        if text.startswith("Abstract:"):
            paragraph.text = text
        if text.startswith("Keywords:"):
            paragraph.text = text

    # 表格内文字做基础统一，避免和正文风格脱节。
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if is_mostly_ascii(run.text):
                            set_run_font(run, "宋体", "Times New Roman", 10.5)
                        else:
                            set_run_font(run, "宋体", "Times New Roman", 10.5)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = apply_template_rules()
    print(path)
