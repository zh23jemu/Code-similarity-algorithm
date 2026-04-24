from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测_格式修订版.docx"


def add_field_run(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = placeholder
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder_text)
    run._r.append(fld_char_end)


def insert_paragraph_before(paragraph, text: str = "", style: str | None = None):
    new_p = paragraph.insert_paragraph_before("")
    if style:
        new_p.style = style
    if text:
        new_p.add_run(text)
    return new_p


def renumber_second_level_titles(doc: Document) -> None:
    """把二级标题统一改成（ 一 ）形式。"""

    chapter_no = 0
    second_no = 0
    cn_nums = "一二三四五六七八九十"

    for p in doc.paragraphs:
        text = p.text.strip()
        if p.style.name == "一级标题":
            chapter_no += 1
            second_no = 0
        elif p.style.name == "二级标题":
            second_no += 1
            title = re.sub(r"^（[一二三四五六七八九十]+）", "", text).strip()
            if second_no <= len(cn_nums):
                p.text = f"（{cn_nums[second_no - 1]}）{title}"


def normalize_keywords(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("关键词："):
            # 保留最核心的四个关键词，满足格式规范
            p.text = "关键词：代码相似度检测；弱监督学习；特征工程；机器学习"
        elif text.startswith("Keywords:"):
            p.text = "Keywords: code similarity detection  weak supervision  feature engineering  machine learning"


def ensure_front_matter(doc: Document) -> None:
    """补前置结构：原创声明、目录。"""

    # 找到中文标题位置，在它之前插入结构页
    title_para = None
    for p in doc.paragraphs:
        if p.text.strip() == "基于弱监督训练建模的本科代码相似度检测算法设计与实现":
            title_para = p
            break
    if title_para is None:
        return

    # 原创声明
    p = insert_paragraph_before(title_para, "", "Normal")
    p.add_run().add_break(WD_BREAK.PAGE)
    p = insert_paragraph_before(title_para, "学位论文（设计）原创声明", "一级标题")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = insert_paragraph_before(
        title_para,
        "本人郑重声明：所呈交的本科毕业论文（设计）是本人在指导教师指导下独立进行研究所取得的成果。除文中已经注明引用的内容外，本论文不包含任何他人已经发表或撰写过的研究成果。",
        "Normal",
    )
    p = insert_paragraph_before(title_para, "作者签名：__________        日期：______年____月____日", "Normal")
    p = insert_paragraph_before(title_para, "", "Normal")
    p.add_run().add_break(WD_BREAK.PAGE)

    # 目录页
    p = insert_paragraph_before(title_para, "目录", "一级标题")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = insert_paragraph_before(title_para, "", "Normal")
    add_field_run(p, r'TOC \o "1-3" \h \z \u', "目录将在 Word 中右键更新域后显示。")
    p = insert_paragraph_before(title_para, "", "Normal")
    p.add_run().add_break(WD_BREAK.PAGE)


def ensure_back_matter(doc: Document) -> None:
    """补致谢和附录，不修改现有正文内容。"""

    # 找到参考文献标题，在其前插入致谢；在其后追加附录
    ref_para = None
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            ref_para = p
            break
    if ref_para is None:
        return

    # 如果不存在致谢，插在参考文献前
    if not any(p.text.strip() == "致谢" for p in doc.paragraphs):
        p = insert_paragraph_before(ref_para, "", "Normal")
        p.add_run().add_break(WD_BREAK.PAGE)
        insert_paragraph_before(ref_para, "致谢", "一级标题")
        insert_paragraph_before(
            ref_para,
            "在论文完成过程中，指导教师、同学及相关资料提供者给予了支持和帮助，在此一并表示感谢。",
            "Normal",
        )

    # 如果不存在附录，则在文末追加
    if not any("附录" == p.text.strip() for p in doc.paragraphs):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_paragraph("附录", style="一级标题")
        doc.add_paragraph("附录可用于放置补充命令、实验输出说明、额外统计表或程序运行截图说明。", style="Normal")


def main() -> None:
    doc = Document(DOC_PATH)
    normalize_keywords(doc)
    renumber_second_level_titles(doc)
    ensure_front_matter(doc)
    ensure_back_matter(doc)
    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()
