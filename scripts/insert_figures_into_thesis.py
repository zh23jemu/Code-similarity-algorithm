from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测_格式修订版.docx"
FIG_DIR = PROJECT_ROOT / "paper_figures"


def insert_after(paragraph, text: str = "", style: str | None = None):
    """在指定段落后插入新段落。"""

    new_p = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_p._p)
    if style:
        new_p.style = style
    if text:
        new_p.add_run(text)
    return new_p


def add_figure_block(anchor_paragraph, image_path: Path, caption: str, source: str, intro_text: str | None = None):
    """在锚点段落后插入 图片+图题+来源。"""

    current = anchor_paragraph
    if intro_text:
        current = insert_after(current, intro_text, "Normal")

    img_p = insert_after(current, style="Normal")
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Cm(14.5))

    caption_p = insert_after(img_p, caption, "Caption")
    insert_after(caption_p, source, "annotation text")


def main() -> None:
    doc = Document(DOC_PATH)

    targets = {p.text.strip(): p for p in doc.paragraphs}

    # 1. 在模型训练与预测流程中加入命令行运行图
    p = targets.get("此外，系统还支持模型对比实验、特征消融实验、严格泛化实验、全量打分结果导出以及人工复核抽样。为了避免人工验证仅集中在最高分样本上，系统在输出高相似代码对之外，还能够导出全部候选代码对的打分结果，并按高分、中间分数、低分三个区间构造分层人工复核样本，为后续更严格的人工验证提供支持。")
    if p is not None:
        add_figure_block(
            p,
            FIG_DIR / "fig_cli_pipeline_run.png",
            "图3. 模型训练与流水线运行命令行输出",
            "数据来源：本系统命令行运行结果整理",
            "为直观展示系统训练与批处理流程，本文在程序实现阶段通过命令行完成数据读取、样本构造、模型训练与结果导出。其典型运行界面如图3所示。",
        )

    # 2. 在实验环境与配置中加入数据概况统计图
    p = targets.get("5. 在严格泛化实验中，默认移除 base_similarity 特征，并按题目划分训练集与测试集；")
    if p is not None:
        add_figure_block(
            p,
            FIG_DIR / "fig_cli_inspect_data.png",
            "图4. 数据集概况统计命令行输出",
            "数据来源：本系统命令行运行结果整理",
            "在正式训练前，系统会先对数据集规模、题目数量和用户数量进行统计，便于确认实验输入的完整性与可用性，如图4所示。",
        )

    # 3. 在模型对比实验部分插入模型对比图和命令行图
    p = targets.get("尽管如此，该实验仍具有一定参考价值。首先，规则基线可作为系统最基础的比较对象；其次，逻辑回归和随机森林表明人工设计特征能够支持传统监督分类建模；最后，对比实验说明在当前特征体系下，随机森林更适合作为后续严格实验的主模型。为了降低“标签规则与特征同源”导致的解释偏差，本文进一步加入了特征消融实验和更严格的泛化实验。")
    if p is not None:
        add_figure_block(
            p,
            FIG_DIR / "fig_model_comparison.png",
            "图5. 模型对比实验结果可视化",
            "数据来源：本系统模型对比实验结果整理",
            "除表格展示外，本文进一步将规则基线、逻辑回归和随机森林的主要指标绘制为柱状图，以更直观地比较不同模型在当前任务中的表现，见图5。",
        )
        # 再插一张命令行图
        anchor = doc.paragraphs[doc.paragraphs.index(p) + 4] if False else p
        add_figure_block(
            p,
            FIG_DIR / "fig_cli_compare_models.png",
            "图6. 模型对比实验命令行运行示意",
            "数据来源：本系统命令行运行结果整理",
            "模型对比实验通过统一命令行入口完成，自动输出训练样本与对比结果文件，典型执行过程如图6所示。",
        )

    # 4. 在消融实验后插入消融图
    p = targets.get("实验结果显示，完整特征集对应的模型表现最好；当去除 base_similarity 后，模型性能仅出现较小幅度下降，说明除基础相似度外，Token 序列、词法统计和结构差异等特征同样能够提供有效判别信息。与此同时，仅使用 base_similarity 也能取得较高结果，说明规则相似度本身已经具备较强表达能力。综合来看，本文所设计的多维特征体系具有一定冗余性和互补性，在弱监督环境下能够保持稳定效果。")
    if p is not None:
        add_figure_block(
            p,
            FIG_DIR / "fig_ablation_f1.png",
            "图7. 特征消融实验 F1 值对比",
            "数据来源：本系统特征消融实验结果整理",
            "为进一步观察不同特征组合对模型性能的影响，本文将消融实验中的 F1 值绘制成图，如图7所示。",
        )

    # 5. 在严格泛化实验后插入严格实验图
    p = targets.get("实验结果表明，在 68 个训练题目和 30 个测试题目条件下，随机森林模型在 48152 个训练样本和 19947 个测试样本上取得了较高性能，其中准确率为 0.99995，精确率为 0.99913，召回率为 1.00000，F1 值为 0.99957，ROC-AUC 为 1.00000。该结果说明，本文设计的特征不仅能够在同题环境下工作，也能够在一定程度上迁移到未参与训练的新题目上，从而增强了方法的泛化可信度。")
    if p is not None:
        add_figure_block(
            p,
            FIG_DIR / "fig_question_split_metrics.png",
            "图8. 按题目划分严格泛化实验结果",
            "数据来源：本系统严格泛化实验结果整理",
            "为了更直观展示按题目划分后的泛化性能，本文进一步将严格实验的主要指标绘制成图，如图8所示。",
        )

    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()
