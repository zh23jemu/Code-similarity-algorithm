from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PROJECT_ROOT / "2.1工科论文参考模板.docx"
SOURCE_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测.docx"
OUTPUT_PATH = PROJECT_ROOT / "毕业论文_代码相似度检测_模板套用版.docx"


def add_field_run(paragraph, instruction: str, placeholder: str = "") -> None:
    """向段落中插入 Word 域。"""

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = placeholder

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder_text)
    run._r.append(fld_char_end)


def iter_block_items(parent):
    """按文档中出现的顺序遍历段落和表格。"""

    if hasattr(parent, "element"):
        parent_element = parent.element.body
    else:
        parent_element = parent._element

    for child in parent_element.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def clear_document_body(doc: Document) -> None:
    """清空模板正文，但保留最后的节属性定义。"""

    body = doc._element.body
    children = list(body)
    for child in children[:-1]:
        body.remove(child)


def add_cover_page(target_doc: Document) -> None:
    """在模板套用版前添加一个可填写封面页。"""

    title = target_doc.add_paragraph(style="摘要、参考文献、注释")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("南京财经大学本科毕业设计（论文）")

    target_doc.add_paragraph("")
    paper_title = target_doc.add_paragraph(style="摘要、参考文献、注释")
    paper_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paper_title.add_run("基于弱监督训练建模的本科代码相似度检测算法设计与实现")

    target_doc.add_paragraph("")
    target_doc.add_paragraph("")

    for line in [
        "学    院：财政与税务学院（请按实际修改）",
        "专    业：__________________________",
        "学生姓名：__________________________",
        "学    号：__________________________",
        "指导教师：__________________________",
        "论文题目：基于弱监督训练建模的本科代码相似度检测算法设计与实现",
        "完成日期：2026年____月____日",
    ]:
        paragraph = target_doc.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)

    target_doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)


def add_toc_page(target_doc: Document) -> None:
    """添加目录页。"""

    heading = target_doc.add_paragraph(style="一级标题")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("目录")

    paragraph = target_doc.add_paragraph(style="Normal")
    add_field_run(paragraph, r'TOC \o "1-3" \h \z \u', "目录将在 Word 中右键更新域后显示。")
    target_doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)


def set_footer_page_number(target_doc: Document) -> None:
    """在页脚中加入页码域。"""

    for section in target_doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        add_field_run(paragraph, "PAGE", "1")


def apply_template_header_footer(target_doc: Document, template_doc: Document) -> None:
    """把模板的页眉信息和页脚样式显式应用到目标文档。

    之前的问题在于：正文重建后，模板原有的节级页眉引用没有稳定保留下来，
    导致目标文档虽然有内容，但页眉显示为空。这里直接把模板页眉文本和
    页脚页码重新写回目标文档，保证最终文档打开时能看到页眉页脚。
    """

    template_header_text = "南京财经大学本科毕业设计"
    if template_doc.sections and template_doc.sections[0].header.paragraphs:
        template_header_text = "".join(p.text for p in template_doc.sections[0].header.paragraphs).strip() or template_header_text

    for section in target_doc.sections:
        section.different_first_page_header_footer = False

        header = section.header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.text = template_header_text


def add_paragraph_from_source(target_doc: Document, paragraph: Paragraph) -> None:
    """把源文档段落按模板样式重建到目标文档。"""

    style_name = paragraph.style.name if paragraph.style is not None else "Normal"
    try:
        new_paragraph = target_doc.add_paragraph(style=style_name)
    except KeyError:
        new_paragraph = target_doc.add_paragraph(style="Normal")

    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline


def add_table_from_source(target_doc: Document, table: Table) -> None:
    """按源表格内容在模板中重建表格。"""

    if not table.rows or not table.columns:
        return

    row_count = len(table.rows)
    col_count = len(table.columns)
    new_table = target_doc.add_table(rows=row_count, cols=col_count)
    try:
        new_table.style = table.style.name if table.style is not None else "Table Grid"
    except Exception:
        new_table.style = "Table Grid"

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            new_table.cell(row_idx, col_idx).text = cell.text


def fill_template_from_source() -> Path:
    """复制模板并用现有论文内容替换正文。

    这里不再直接拷贝源文档 XML，而是逐段落、逐表格重建到模板里。
    这样新的段落会真正使用模板内部已有的样式定义，格式会比直接搬运 XML
    更接近学校模板。
    """

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"找不到模板文件: {TEMPLATE_PATH}")
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"找不到源论文文件: {SOURCE_PATH}")

    shutil.copyfile(TEMPLATE_PATH, OUTPUT_PATH)

    source_doc = Document(SOURCE_PATH)
    template_doc = Document(TEMPLATE_PATH)
    target_doc = Document(OUTPUT_PATH)
    clear_document_body(target_doc)

    add_cover_page(target_doc)
    add_toc_page(target_doc)

    for block in iter_block_items(source_doc):
        if isinstance(block, Paragraph):
            add_paragraph_from_source(target_doc, block)
        elif isinstance(block, Table):
            add_table_from_source(target_doc, block)

    apply_template_header_footer(target_doc, template_doc)
    set_footer_page_number(target_doc)

    target_doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output_path = fill_template_from_source()
    print(output_path)
